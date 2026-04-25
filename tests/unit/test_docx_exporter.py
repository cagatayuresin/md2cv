"""DOCXExporter tests — covers every parsing branch including table."""

import zipfile
from io import BytesIO
from pathlib import Path

from docx import Document

from app.core.exporters import DOCXExporter


def _open(data: bytes):
    return Document(BytesIO(data))


def test_export_bytes_is_valid_zip_docx(minimal_md: str) -> None:
    parser_meta = {"name": "Jane Doe"}
    body = "# Section\n\n- item one\n- item two\n"
    data = DOCXExporter().export_bytes(parser_meta, body)
    assert zipfile.is_zipfile(BytesIO(data))


def test_export_to_path_writes_file(tmp_path: Path) -> None:
    out = tmp_path / "cv.docx"
    DOCXExporter().export_to_path({"name": "Jane"}, "# Hi\n", out)
    assert out.exists()
    assert zipfile.is_zipfile(out)


def test_header_includes_name_title_contact_links() -> None:
    meta = {
        "name": "Jane Doe",
        "title": "Engineer",
        "email": "j@x.com",
        "phone": "+1",
        "location": "Remote",
        "linkedin": "linkedin.com/in/jane",
        "github": "github.com/jane",
        "website": "jane.dev",
    }
    data = DOCXExporter().export_bytes(meta, "# Body\n")
    doc = _open(data)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Jane Doe" in text
    assert "Engineer" in text
    assert "j@x.com" in text
    assert "linkedin.com/in/jane" in text


def test_headings_h1_h2_h3() -> None:
    body = "# H1\n\n## H2\n\n### H3\n"
    data = DOCXExporter().export_bytes({}, body)
    doc = _open(data)
    styles = [p.style.name for p in doc.paragraphs]
    assert "Heading 1" in styles
    assert "Heading 2" in styles


def test_bullet_list_creates_list_bullet_style() -> None:
    body = "# X\n\n- one\n- two\n- three\n"
    data = DOCXExporter().export_bytes({}, body)
    doc = _open(data)
    styles = [p.style.name for p in doc.paragraphs]
    assert any(s == "List Bullet" for s in styles)


def test_star_bullet_list_also_works() -> None:
    body = "# X\n\n* alpha\n* beta\n"
    data = DOCXExporter().export_bytes({}, body)
    doc = _open(data)
    texts = [p.text for p in doc.paragraphs]
    assert "alpha" in texts


def test_horizontal_rule_and_html_comments_skipped() -> None:
    body = "# H\n\n---\n\n<!-- comment -->\n\nReal text.\n"
    data = DOCXExporter().export_bytes({}, body)
    doc = _open(data)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Real text." in text
    assert "comment" not in text


def test_standalone_bold_paragraph_is_bold() -> None:
    body = "# X\n\n**Important**\n"
    data = DOCXExporter().export_bytes({}, body)
    doc = _open(data)
    bold_paragraphs = [p for p in doc.paragraphs if any(r.bold for r in p.runs)]
    assert any("Important" in p.text for p in bold_paragraphs)


def test_links_are_stripped_to_text() -> None:
    body = "# X\n\n[website](https://example.com)\n"
    data = DOCXExporter().export_bytes({}, body)
    doc = _open(data)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "website" in text
    assert "https://example.com" not in text


def test_table_round_trip() -> None:
    body = "# Skills\n\n| Area | Tools |\n|------|-------|\n| Lang | Python |\n| Cloud | AWS |\n"
    data = DOCXExporter().export_bytes({}, body)
    doc = _open(data)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    rows = [[c.text for c in r.cells] for r in table.rows]
    assert rows[0] == ["Area", "Tools"]
    assert ["Lang", "Python"] in rows
    assert ["Cloud", "AWS"] in rows


def test_clean_markdown_removes_html_and_markers() -> None:
    cleaned = DOCXExporter._clean_markdown("**Bold** and *italic* with <span>html</span>")
    assert cleaned == "Bold and italic with html"


def test_h3_is_italic_paragraph() -> None:
    data = DOCXExporter().export_bytes({}, "### Subheading\n")
    doc = _open(data)
    italic = [p for p in doc.paragraphs if any(r.italic for r in p.runs)]
    assert any("Subheading" in p.text for p in italic)
