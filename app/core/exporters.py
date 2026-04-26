"""PDF and DOCX exporters."""

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from weasyprint import CSS, HTML

from app.core.renderer import TemplateRenderer


class PDFExporter:
    """Export rendered HTML to PDF using WeasyPrint."""

    def __init__(self, renderer: TemplateRenderer):
        self.renderer = renderer

    def _stylesheets(self) -> list[CSS]:
        css_path = self.renderer.get_css_path()
        if css_path.exists():
            return [CSS(filename=str(css_path))]
        return []

    def export_bytes(self, html_content: str) -> bytes:
        html = HTML(string=html_content, base_url=str(self.renderer.template_dir))
        result = html.write_pdf(stylesheets=self._stylesheets())
        return bytes(result) if result is not None else b""

    def export_to_path(self, html_content: str, output_path: Path) -> Path:
        output_path = Path(output_path)
        html = HTML(string=html_content, base_url=str(self.renderer.template_dir))
        html.write_pdf(str(output_path), stylesheets=self._stylesheets())
        return output_path


class DOCXExporter:
    """Export CV content to DOCX format."""

    HORIZONTAL_RULE_LENGTH = 60
    TABLE_SEPARATOR_PATTERN = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self) -> None:
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Georgia"
        normal.font.size = Pt(11)

        h1 = styles["Heading 1"]
        h1.font.name = "Arial"
        h1.font.size = Pt(14)
        h1.font.bold = True

        h2 = styles["Heading 2"]
        h2.font.name = "Arial"
        h2.font.size = Pt(12)
        h2.font.bold = True

    def export_bytes(self, meta: dict, markdown_body: str) -> bytes:
        self._build(meta, markdown_body)
        buffer = BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()

    def export_to_path(self, meta: dict, markdown_body: str, output_path: Path) -> Path:
        output_path = Path(output_path)
        self._build(meta, markdown_body)
        self.doc.save(str(output_path))
        return output_path

    def _build(self, meta: dict, markdown_body: str) -> None:
        self._add_header(meta)
        self._add_markdown_content(markdown_body)

    def _add_header(self, meta: dict) -> None:
        if meta.get("name"):
            title = self.doc.add_heading(meta["name"], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if meta.get("title"):
            subtitle = self.doc.add_paragraph(meta["title"])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if subtitle.runs:
                subtitle.runs[0].italic = True

        contact_parts = [meta[k] for k in ("email", "phone", "location") if meta.get(k)]
        if contact_parts:
            contact = self.doc.add_paragraph(" | ".join(contact_parts))
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        link_parts = [meta[k] for k in ("linkedin", "github", "website") if meta.get(k)]
        if link_parts:
            links = self.doc.add_paragraph(" | ".join(link_parts))
            links.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph("_" * self.HORIZONTAL_RULE_LENGTH)

    def _add_markdown_content(self, markdown_body: str) -> None:
        lines = markdown_body.strip().split("\n")
        i = 0
        current_list: list[str] = []

        while i < len(lines):
            stripped = lines[i].strip()

            if not stripped:
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                i += 1
                continue

            if stripped.startswith("---") or stripped.startswith("<!--"):
                i += 1
                continue

            if self._is_table_start(lines, i):
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                consumed = self._add_table(lines, i)
                i += consumed
                continue

            if stripped.startswith("# "):
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                self.doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                self.doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                paragraph = self.doc.add_paragraph(stripped[4:])
                if paragraph.runs:
                    paragraph.runs[0].italic = True
            elif stripped.startswith("- ") or stripped.startswith("* "):
                current_list.append(stripped[2:])
            elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run(stripped[2:-2])
                run.bold = True
            else:
                if current_list:
                    self._add_list(current_list)
                    current_list = []
                text = self._clean_markdown(stripped)
                if text:
                    self.doc.add_paragraph(text)

            i += 1

        if current_list:
            self._add_list(current_list)

    def _is_table_start(self, lines: list[str], index: int) -> bool:
        if index + 1 >= len(lines):
            return False
        header = lines[index].strip()
        separator = lines[index + 1].strip()
        if not (header.startswith("|") and header.endswith("|")):
            return False
        return bool(self.TABLE_SEPARATOR_PATTERN.match(separator))

    def _add_table(self, lines: list[str], start: int) -> int:
        header_cells = self._split_table_row(lines[start])
        rows: list[list[str]] = []
        i = start + 2
        while i < len(lines):
            stripped = lines[i].strip()
            if not stripped or not stripped.startswith("|"):
                break
            rows.append(self._split_table_row(lines[i]))
            i += 1

        cols = len(header_cells)
        table = self.doc.add_table(rows=1 + len(rows), cols=cols)
        table.style = "Table Grid"

        header_row = table.rows[0]
        for j, cell_text in enumerate(header_cells):
            cell = header_row.cells[j]
            cell.text = self._clean_markdown(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for row_idx, row_cells in enumerate(rows, start=1):
            row = table.rows[row_idx]
            for j in range(cols):
                value = row_cells[j] if j < len(row_cells) else ""
                row.cells[j].text = self._clean_markdown(value)

        self.doc.add_paragraph()
        return (i - start)

    @staticmethod
    def _split_table_row(line: str) -> list[str]:
        stripped = line.strip()
        if stripped.startswith("|"):
            stripped = stripped[1:]
        if stripped.endswith("|"):
            stripped = stripped[:-1]
        return [cell.strip() for cell in stripped.split("|")]

    def _add_list(self, items: list[str]) -> None:
        for item in items:
            text = self._clean_markdown(item)
            self.doc.add_paragraph(text, style="List Bullet")

    @staticmethod
    def _clean_markdown(text: str) -> str:
        # Bounded character classes (no `.+?` with overlapping delimiters) so
        # these substitutions can't blow up via catastrophic backtracking on
        # adversarial input (S5852).
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
        text = re.sub(r"<[^>]+>", "", text)
        return text.strip()
